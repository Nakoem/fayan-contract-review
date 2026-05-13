"""
生成设计感简历docx — 深蓝+金色点缀风格
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Pt, RGBColor

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10)

NAVY = RGBColor(26, 31, 54)
GOLD = RGBColor(201, 169, 110)
DARK = RGBColor(44, 36, 22)
GRAY = RGBColor(92, 82, 64)
WHITE = RGBColor(255, 255, 255)


def add_section(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("▌ ")
    run.font.size = Pt(14)
    run.font.name = "微软雅黑"
    run.font.color.rgb = GOLD
    run.bold = True
    run = p.add_run(f" {text}")
    run.font.size = Pt(13)
    run.font.name = "微软雅黑"
    run.font.color.rgb = NAVY
    run.bold = True
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f"<w:pBdr {nsdecls('w')}>"
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="DFD5B8"/>'
        f"</w:pBdr>"
    )
    pPr.append(pBdr)


def add_bullet(text, indent=0.4, size=10, bold_prefix="", color=DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run("• ")
    run.font.size = Pt(size)
    run.font.name = "微软雅黑"
    run.font.color.rgb = GOLD
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(size)
        run.font.name = "微软雅黑"
        run.font.color.rgb = DARK
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "微软雅黑"
    run.font.color.rgb = color


def add_job(company, title, period, items):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(company)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = "微软雅黑"
    run.font.color.rgb = NAVY
    run = p.add_run(f"    {title}    ")
    run.font.size = Pt(9.5)
    run.font.name = "微软雅黑"
    run.font.color.rgb = GRAY
    run = p.add_run(period)
    run.font.size = Pt(9)
    run.font.name = "微软雅黑"
    run.font.color.rgb = GRAY
    for item in items:
        add_bullet(item, indent=0.4, size=9.5)


# ═══════════════════════════════════════
# 头部
# ═══════════════════════════════════════
header = doc.add_table(rows=1, cols=1)
tbl = header._tbl
tblPr = parse_xml(f"<w:tblPr {nsdecls('w')}/>")
tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="dxa"/>')
tblPr.append(tblW)

cell = header.rows[0].cells[0]
shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A1F36"/>')
cell._tc.get_or_add_tcPr().append(shading)

p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
run = p.add_run("蒲 冠 宇")
run.font.size = Pt(24)
run.font.name = "微软雅黑"
run.font.color.rgb = WHITE
run.bold = True

p = cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
run = p.add_run("AI 应 用 工 程 师")
run.font.size = Pt(12)
run.font.name = "微软雅黑"
run.font.color.rgb = GOLD

p = cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(14)
run = p.add_run("18984961952  |  627101067@qq.com  |  github.com/pppppgy  |  男 · 31岁")
run.font.size = Pt(9)
run.font.name = "微软雅黑"
run.font.color.rgb = RGBColor(180, 192, 208)

doc.add_paragraph()

# ═══════════════════════════════════════
# 个人概述
# ═══════════════════════════════════════
add_section("个人概述")

p = doc.add_paragraph()
run = p.add_run(
    "3年审计与风控经验。因工作中频繁接触合同审查，发现AI替代人工审查的潜力，"
    '自学Python和大模型应用开发，4个月内从零独立交付了"法眼"——产品级AI合同审查Agent。'
    '信奉"够用就好"的实用主义：技术服务于业务，不造轮子。日常使用Claude Code辅助开发。'
)
run.font.size = Pt(10)
run.font.name = "微软雅黑"
run.font.color.rgb = DARK

# ═══════════════════════════════════════
# AI项目
# ═══════════════════════════════════════
add_section("AI 项目经历")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(2)
run = p.add_run("法眼 · AI 合同审查系统")
run.bold = True
run.font.size = Pt(11.5)
run.font.name = "微软雅黑"
run.font.color.rgb = NAVY
run = p.add_run("  —  独立全栈开发 · 4个月")
run.font.size = Pt(9.5)
run.font.name = "微软雅黑"
run.font.color.rgb = GRAY

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run(
    'AI Agent替代人工"读合同→查法规→写风控报告"。6种合同类型，66条法规/判例/政策/税务知识→RAG语义检索。'
    "提供Web/CLI/API/MCP四种使用方式。  GitHub：github.com/pppppgy/fayan-contract-review"
)
run.font.size = Pt(9.5)
run.font.name = "微软雅黑"
run.font.color.rgb = GRAY

items1 = [
    (
        "Agent架构：",
        '设计10工具ReAct自主决策循环，自动完成"提取条款→查法规→逐条分析→完整性检查→生成报告"，每种合同配备法定红线标准。',
    ),
    (
        "RAG全链路：",
        '66条知识→Chroma向量库（216chunks），DashScope Embedding语义检索。效果："押金不退"命中"押金退还"（关键词0%→语义91%）。',
    ),
    (
        "模型选型：",
        "对比测试4个模型（qwen-plus/max/3.6-plus/deepseek-v3.2），定义4维评测标准，选qwen-plus。",
    ),
    (
        "工程化：",
        "FastAPI REST + Docker + loguru日志 + Prompt YAML版管 + 三维度自动化评估（格式/召回率/LLM裁判）。",
    ),
    ("MCP协议：", "5个法律搜索工具封装为MCP Server，Claude Code/Cursor可直接调用。"),
]
for bp, text in items1:
    add_bullet(text, indent=0.4, bold_prefix=bp)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.4)
p.paragraph_format.space_before = Pt(4)
run = p.add_run(
    "Python  ·  FastAPI  ·  Streamlit  ·  Docker  ·  Chroma  ·  MCP SDK  ·  DashScope  ·  Claude Code"
)
run.font.size = Pt(9)
run.font.name = "微软雅黑"
run.font.color.rgb = GOLD

# 项目2
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(2)
run = p.add_run("法眼 · 法律问答 Bot")
run.bold = True
run.font.size = Pt(11.5)
run.font.name = "微软雅黑"
run.font.color.rgb = NAVY
run = p.add_run("  —  1周")
run.font.size = Pt(9.5)
run.font.name = "微软雅黑"
run.font.color.rgb = GRAY

add_bullet(
    "跨5库RAG检索 + SSE流式输出 + 20轮对话记忆 + 法条出处标注。双模式：不传合同→泛法律问答；上传合同→针对条款定向问答。",
    indent=0.4,
)

# ═══════════════════════════════════════
# 工作经历
# ═══════════════════════════════════════
add_section("工作经历")

jobs = [
    (
        "贵州创壹佳信息服务有限公司",
        "风控专员",
        "2025.01 - 2026.03",
        [
            '审核客户资质及信用资料，按风控标准出具风险评估意见——与AI审查的"评估→出具意见"流程同构',
            "跟踪履约情况动态调整风险等级，参与完善风控制度和流程优化",
            "分析业务数据形成风控报告——培养数据驱动决策的思维习惯",
        ],
    ),
    (
        "遵义中审会计师事务所",
        "审计助理",
        "2023.12 - 2024.12",
        [
            "运用审计软件编制工作底稿、凭证抽查、账务核对——逐条核查的工作方式与AI分析的精确性一致",
            "检查记账凭证发现遗漏信息，完善审计证据链——对细节和证据完整性高度敏感",
            "直接与客户财务沟通，准确理解业务需求——3年客户沟通经验",
        ],
    ),
    (
        "中信证券华南股份有限公司",
        "证券事务专员",
        "2020.06 - 2023.01",
        [
            "客户账户管理及日常证券交易业务咨询——面对客户的答疑训练了快速理解与解决问题能力",
            "组织投资者教育活动，提供证券市场培训和政策解读——向非专业人士解释复杂规则",
            "严格执行合规风控制度——合规意识深入工作习惯",
        ],
    ),
]

for job in jobs:
    add_job(*job)

# ═══════════════════════════════════════
# 核心技能
# ═══════════════════════════════════════
add_section("核心技能")

skills = [
    (
        "AI/Agent",
        "ReAct Agent · Function Calling · RAG语义检索 · Prompt工程 · MCP协议 · LLM-as-Judge",
    ),
    ("后端/部署", "Python · FastAPI · REST API · Docker容器化 · docker-compose"),
    ("数据/检索", "Chroma向量数据库 · DashScope Embedding · 混合检索（语义+关键词）"),
    ("工程化", "loguru日志系统 · YAML配置管理 · Prompt版本管理 · 自动化评估"),
    ("工具链", "Claude Code（日常主力开发）· OpenAI兼容SDK · Streamlit · Git/GitHub"),
    ("领域知识", "审计准则 · 企业会计准则 · 风控流程 · 合同审查（3年实务经验）"),
]

for cat, items in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(cat)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "微软雅黑"
    run.font.color.rgb = NAVY
    run = p.add_run(f"  {items}")
    run.font.size = Pt(10)
    run.font.name = "微软雅黑"
    run.font.color.rgb = DARK

# ═══════════════════════════════════════
# 教育
# ═══════════════════════════════════════
add_section("教育经历")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
run = p.add_run("暨南大学  ·  护理学  ·  本科")
run.font.size = Pt(10.5)
run.font.name = "微软雅黑"
run.font.color.rgb = NAVY

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1)
run = p.add_run("自学Python开发、大模型应用、AI Agent架构  |  大学英语6级  |  证券从业资格")
run.font.size = Pt(9.5)
run.font.name = "微软雅黑"
run.font.color.rgb = GRAY

# ── 保存 ──
output = Path.home() / "Desktop" / "个人简历-AI应用工程师-设计版.docx"
doc.save(str(output))
print(f"Saved to: {output}")
